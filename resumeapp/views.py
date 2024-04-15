import io
import os
import re
import zipfile
import tempfile
import xlsxwriter

import pandas as pd
from django.shortcuts import render, redirect
from django.http import HttpResponse, HttpResponseBadRequest
from django.views.decorators.csrf import csrf_exempt
from .forms import ResumeUploadForm
from PyPDF2 import PdfReader
from docx import Document

# Helper functions
def extract_text_and_info(file):
    data = {'email': [], 'contact': [], 'text': []}
    if file.endswith('.pdf'):
        text = extract_text_from_pdf(file)
    elif file.endswith('.docx'):
        text = extract_text_from_docx(file)
    elif file.endswith('.doc'):
        text = extract_text_from_doc(file)
    else:
        return data

    # Extract email addresses
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    data['email'].extend(emails)

    # Extract contact numbers
    contact_pattern = r'(\+\d{10,}[-\.\s]??\d{1,4}[-\.\s]??\d{1,4}[-\.\s]??\d{1,4})|(\(\d{2,6}\)\s*\d{1,5}[-\.\s]*\d{1,5}[-\.\s]*\d{1,4})|(\d{10,}[-\.\s]\d{1,5}[-\.\s]\d{1,5})'
    contacts = re.findall(contact_pattern, text)
    data['contact'].extend([match[0] or match[1] or match[2] for match in contacts])

    data['text'].append(text)

    return data

def extract_text_from_pdf(pdf_file):
    with tempfile.NamedTemporaryFile(delete=False) as f:
        f.write(pdf_file.read())
        f.seek(0)
        pdf_reader = PdfReader(f)
        text = ''
        for page in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page].extract_text()
    return text

def extract_text_from_docx(docx_file):
    with tempfile.NamedTemporaryFile(delete=False) as f:
        f.write(docx_file.read())
        f.seek(0)
        document = Document(f)
        text = ' '.join([paragraph.text for paragraph in document.paragraphs])
    return text

def extract_text_from_doc(doc_file):
    with tempfile.NamedTemporaryFile(delete=False) as f:
        f.write(doc_file.read())
        f.seek(0)
        doc = Document(f)
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text
    return text

def extract_email(text):
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    return re.findall(email_pattern, text)

def extract_contact(text):
    contact_pattern = r'(\+\d{1,3}[-\.\s]??\d{1,4}[-\.\s]??\d{1,4}[-\.\s]??\d{1,4})|(\(\d{2,6}\)\s*\d{1,5}[-\.\s]*\d{1,5}[-\.\s]*\d{1,5})|(\d{3,5}[-\.\s]\d{1,5}[-\.\s]\d{1,5})'
    matches = re.findall(contact_pattern, text)
    return [match[0] or match[1] or match[2] for match in matches]

def download_report(output, filename):
    response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def process_file(file):
    data = {'email': [], 'contact': [], 'text': []}
    zip_file = zipfile.ZipFile(file)

    with tempfile.TemporaryDirectory() as temp_dir:
        for filename in zip_file.namelist():
            print(filename)
            try:
                if filename.endswith('.pdf'):
                    with zip_file.open(filename) as f:
                        text = extract_text_from_pdf(f)
                elif filename.endswith('.docx'):
                    with zip_file.open(filename) as f:
                        text = extract_text_from_docx(f)
                elif filename.endswith('.doc'):
                    with zip_file.open(filename) as f:
                        text = extract_text_from_doc(f)
                else:
                    continue

                email = extract_email(text)
                contact = extract_contact(text)

                data['email'].append(email)
                data['contact'].append(contact)
                data['text'].append(' '.join(text.split()))

            except zipfile.BadZipFile:
                print(f"Error: {filename} is not a valid zip file")

        df = pd.DataFrame(data)
        output = io.BytesIO()

        workbook = xlsxwriter.Workbook(output)
        worksheet = workbook.add_worksheet()

        # Write column names
        worksheet.write_row(0, 0, ['emailID', 'contact no.', 'text'])

        # Write data, cleaning up formatting
        for row_num, row in enumerate(df.iterrows()):
            email = row[1]['email']
            contact = row[1]['contact']
            text = row[1]['text']

            # Clean up email formatting
            email = re.sub(r'\s+', '', email[0]) if email else ''

            # Clean up contact formatting
            contact = re.sub(r'\D+', '', contact[0]) if contact else ''

            # Write row
            worksheet.write_row(row_num + 1, 0, [email, contact, text])

        workbook.close()
        output.seek(0)

    return download_report(output, 'resumes_report.xlsx')

def upload_resume(request):
    if request.method == 'POST':
        if 'file' in request.FILES:
            return process_file(request.FILES['file'])
        else:
            # Handle the case where the file key is not present in the request.FILES dictionary
            return HttpResponseBadRequest("No file was included in the request.")
    else:
        form = ResumeUploadForm()

    return render(request, 'C:/Users/Priyanshu/Desktop/Assg/resume_extractor/resumeapp/templates/upload_resume.html', {'form': form})

def home(request):
    return render(request, 'C:/Users/Priyanshu/Desktop/Assg/resume_extractor/resumeapp/templates/home.html')